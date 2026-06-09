# =============================================================================
# part4_automated_report/01_report_docx.py
# Génération automatisée d'un rapport Word avec python-docx
# Injecte les cartes PNG et les statistiques dans un document structuré
# Prérequis : avoir exécuté les scripts de part1 pour générer les cartes PNG
# =============================================================================

from pathlib import Path
import sys
from datetime import date

import geopandas as gpd
import numpy as np
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

sys.path.insert(0, str(Path(__file__).parent.parent))
from config_loader import load_config, get_vector_dir, get_output_dir


def set_cell_background(cell, hex_color: str) -> None:
    """
    Applique une couleur de fond à une cellule de tableau Word.
    Nécessite une manipulation directe du XML OOXML (python-docx ne le fait pas nativement).

    Paramètres
    ----------
    cell      : cellule docx cible
    hex_color : couleur en hexadécimal sans # (ex. '1a237e')
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def add_section_title(doc: Document, title: str, level: int = 1) -> None:
    """
    Ajoute un titre de section avec le style Heading correspondant.

    Paramètres
    ----------
    doc   : document Word cible
    title : texte du titre
    level : niveau de titre (1, 2 ou 3)
    """
    heading = doc.add_heading(title, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT


def add_figure(doc: Document, image_path: Path, caption: str,
               width_inches: float = 5.5) -> None:
    """
    Insère une figure (image PNG) avec sa légende dans le document.
    La figure est centrée et la légende est en italique sous l'image.

    Paramètres
    ----------
    doc          : document Word cible
    image_path   : chemin vers l'image PNG
    caption      : texte de légende
    width_inches : largeur d'affichage en pouces
    """
    if not image_path.exists():
        doc.add_paragraph(f"[Image manquante : {image_path.name}]").italic = True
        return

    # Paragraphe centré pour l'image
    para_img = doc.add_paragraph()
    para_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_img = para_img.add_run()
    run_img.add_picture(str(image_path), width=Inches(width_inches))

    # Légende en italique, centrée
    para_cap = doc.add_paragraph(caption)
    para_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_cap = para_cap.runs[0]
    run_cap.italic = True
    run_cap.font.size = Pt(9)
    run_cap.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # Espace après la figure
    doc.add_paragraph()


def add_stats_table(doc: Document, gdf: gpd.GeoDataFrame,
                     columns: list, headers: list) -> None:
    """
    Insère un tableau de statistiques descriptives dans le document.
    L'en-tête est coloré et les lignes alternent deux tons de gris.

    Paramètres
    ----------
    doc     : document Word cible
    gdf     : GeoDataFrame source des données
    columns : colonnes à inclure dans le tableau
    headers : étiquettes d'en-tête correspondantes
    """
    n_rows = len(gdf) + 1   # +1 pour l'en-tête
    n_cols = len(columns)

    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # --- Ligne d'en-tête ---
    header_row = table.rows[0]
    for i, header in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = header
        set_cell_background(cell, "1a237e")   # bleu institutionnel
        run = cell.paragraphs[0].runs[0]
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.bold = True
        run.font.size = Pt(9)

    # --- Lignes de données ---
    for row_idx, (_, data_row) in enumerate(gdf.iterrows()):
        row = table.rows[row_idx + 1]
        bg_color = "f5f5f5" if row_idx % 2 == 0 else "ffffff"

        for col_idx, col in enumerate(columns):
            cell = row.cells[col_idx]
            value = data_row[col]

            # Formatage selon le type de valeur
            if isinstance(value, float):
                cell.text = f"{value:,.2f}"
            elif isinstance(value, (int, np.integer)):
                cell.text = f"{value:,}"
            else:
                cell.text = str(value)

            set_cell_background(cell, bg_color)
            cell.paragraphs[0].runs[0].font.size = Pt(9)


def build_report(doc: Document, gdf: gpd.GeoDataFrame,
                  output_dir: Path) -> None:
    """
    Construit le rapport complet en assemblant toutes les sections.

    Paramètres
    ----------
    doc        : document Word vide ou depuis un gabarit
    gdf        : GeoDataFrame des arrondissements avec toutes les variables
    output_dir : répertoire contenant les cartes PNG générées
    """
    # --- Page de titre ---
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run("Rapport de visualisation cartographique\nGMQ 580 — Géo-informatique II")
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)

    doc.add_paragraph()

    date_para = doc.add_paragraph(f"Date de production : {date.today().strftime('%d %B %Y')}")
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para.runs[0].font.size = Pt(11)
    date_para.runs[0].italic = True

    doc.add_page_break()

    # --- Section 1 : Introduction ---
    add_section_title(doc, "1. Introduction", level=1)
    doc.add_paragraph(
        "Ce rapport présente les résultats de l'analyse cartographique des arrondissements "
        "de Sherbrooke. Les données proviennent de la Ville de Sherbrooke et ont été traitées "
        "avec Python (GeoPandas, Matplotlib, mapclassify) dans le cadre du cours GMQ 580."
    )
    doc.add_paragraph(
        f"L'analyse porte sur {len(gdf)} arrondissements. "
        "La variable cartographiée est la superficie calculée à partir des géométries des polygones. "
        "Toutes les superficies sont calculées dans le système de référence EPSG:32198 (MTM Québec)."
    )

    # --- Section 2 : Cartographie thématique ---
    add_section_title(doc, "2. Cartographie thématique", level=1)

    add_section_title(doc, "2.1 Superficie des arrondissements", level=2)
    doc.add_paragraph(
        "La carte suivante présente la superficie des arrondissements (km²), "
        "discrétisée selon la méthode des ruptures naturelles (NaturalBreaks, k=5)."
    )
    add_figure(
        doc,
        image_path=output_dir / "choropleth_natural_breaks.png",
        caption="Figure 1 — Superficie par arrondissement (km²), NaturalBreaks k=5. "
                "Source : Ville de Sherbrooke | Projection : EPSG:32198",
    )

    add_section_title(doc, "2.2 Comparaison des méthodes de discrétisation", level=2)
    doc.add_paragraph(
        "Le choix de la méthode de discrétisation influence significativement la lecture "
        "du phénomène cartographié. La figure ci-dessous compare trois approches courantes."
    )
    add_figure(
        doc,
        image_path=output_dir / "choropleth_methods_comparison.png",
        caption="Figure 2 — Comparaison des méthodes NaturalBreaks, Quantiles et EqualInterval. "
                "Source : Ville de Sherbrooke | Projection : EPSG:32198",
        width_inches=6.0,
    )

    add_section_title(doc, "2.3 Fond de carte contextuel", level=2)
    doc.add_paragraph(
        "L'ajout d'un fond de carte (contextily) enrichit la lecture spatiale "
        "en situant les arrondissements dans leur environnement urbain."
    )
    add_figure(
        doc,
        image_path=output_dir / "basemap_cartodb_choropleth.png",
        caption="Figure 3 — Population par arrondissement sur fond CartoDB Positron. "
                "Source : Ville de Sherbrooke, OpenStreetMap contributors",
    )

    # --- Section 3 : Analyse raster ---
    add_section_title(doc, "3. Analyse raster — NDVI", level=1)
    doc.add_paragraph(
        "L'indice de végétation normalisé (NDVI) est calculé à partir des bandes "
        "rouge (B4) et proche-infrarouge (B5) de l'image Landsat 8/9. "
        "Les valeurs positives élevées indiquent une végétation dense et saine."
    )
    add_figure(
        doc,
        image_path=output_dir / "raster_ndvi.png",
        caption="Figure 4 — NDVI calculé à partir de l'image Landsat 8/9 (Sherbrooke). "
                "Palette RdYlGn : rouge = végétation absente, vert = végétation dense.",
    )

    # --- Section 4 : Tableau de données ---
    add_section_title(doc, "4. Données par arrondissement", level=1)
    doc.add_paragraph(
        "Le tableau suivant présente les attributs géométriques de chaque arrondissement."
    )

    add_stats_table(
        doc=doc,
        gdf=gdf,
        columns=["NOM", "NUMERO", "area_km2"],
        headers=["Arrondissement", "Numéro", "Superficie (km²)"],
    )

    doc.add_paragraph()

    # --- Section 5 : Conclusion ---
    add_section_title(doc, "5. Conclusion", level=1)
    doc.add_paragraph(
        "L'automatisation de la production cartographique avec Python permet de générer "
        "des rapports reproductibles et cohérents. Les scripts sont versionnés (Git) et "
        "peuvent être exécutés à nouveau dès que les données source sont mises à jour, "
        "sans intervention manuelle sur la mise en page."
    )


def main():
    config = load_config()
    vector_dir  = get_vector_dir(config)
    output_dir  = get_output_dir(config)

    crs_projected = config["crs"]["projected"]

    # Chargement et préparation des données
    gdf = gpd.read_file(vector_dir / "Arrondissement_-5179909302249146713.gpkg").to_crs(crs_projected)
    gdf["area_km2"] = gdf.geometry.area / 1_000_000

    # Création d'un document Word vide avec marges standard (2 cm)
    doc = Document()

    # Marges : 2 cm partout (exigence du plan de cours)
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2)
        section.right_margin  = Cm(2)

    # Style de police par défaut : Times New Roman 11pt
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)

    # Construction du rapport
    build_report(doc, gdf, output_dir)

    # Sauvegarde
    output_path = output_dir / "rapport_visualisation_GMQ580.docx"
    doc.save(str(output_path))
    print(f"Rapport sauvegardé : {output_path}")


if __name__ == "__main__":
    main()
